#!/usr/bin/env python3
"""SANDBOX one-off: build the dry-run outcomes sheet as a .docx companion to
docs/calibration_prep/06_DRYRUN_OUTCOMES.md.

Pure python-docx (no pandoc / no anthropic-docx-skill pack.py, so the cp1252
numbering.xml gotcha does not apply). Output is byte-faithful to the .md content,
re-cast for Word and audited against the 12 document-quality rules:
  * R5  emphasis-caps dropped (Sandbox/Retain/Locked/Gated/Structural title-case+bold;
        only true acronyms + system enum values stay upper)
  * R6/R12 column widths sized so atomic values + headers never wrap (nbsp on currency)
  * R8  >=1% deltas rounded to nearest 0.5 (tail +20/+38/+45.5; warm ~3.5)
  * R9  round labels short-form (R5m/R150k); computed aggregates R 124.0m; value cells
        full-form R 750,247 with non-breaking spaces
  * R10 third-person only (asserted by a we/us/our/I/you/your/they/their grep at the end)
R11 (VAT/commission basis) does NOT trigger: modelling/parameter doc, not a rate quote.
NOT shipped (FIN-9 calibration prep, 2026-06-03). Run: py tooling/_build_dryrun_outcomes_docx.py
"""
import os, re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SEC = os.path.dirname(HERE)
OUT = os.path.join(SEC, "docs", "calibration_prep", "06_DRYRUN_OUTCOMES.docx")

NB = " "          # non-breaking space (prevents R / value mid-wrap, R6)
def Rfull(n):          # full-form value cell: R 1,793,092
    return f"R{NB}{n:,}"

HEADER_FILL = "1F3864"   # dark blue
HEADER_TXT = "FFFFFF"
ALT_FILL = "EEF1F7"      # light row banding
ACCENT = RGBColor(0x1F, 0x38, 0x64)

doc = Document()

# --- page: A4 portrait, 16mm margins -> 178mm usable -------------------------
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
for m in ("left_margin", "right_margin"):
    setattr(sec, m, Mm(16))
sec.top_margin, sec.bottom_margin = Mm(15), Mm(15)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.08


def _set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _fixed_layout(table):
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
    trPr.append(th)


def _cell_text(cell, text, size, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    if align is not None:
        p.alignment = align
    # bold the first run if text contains a leading ** marker convention: skip; plain
    run = p.add_run(text)
    run.font.size = Pt(size); run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = 1  # center


def add_title(text, sub):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.keep_with_next = True
    p2 = doc.add_paragraph()
    r2 = p2.add_run(sub); r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p2.paragraph_format.space_after = Pt(6); p2.paragraph_format.keep_with_next = True


def add_meta(label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rl = p.add_run(label + "  "); rl.bold = True; rl.font.size = Pt(9.5); rl.font.color.rgb = ACCENT
    rb = p.add_run(body); rb.font.size = Pt(9.5)


def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True


def add_h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True


def add_para(runs, size=10.5, space_after=4, italic=False, color=None):
    """runs: list of (text, bold) or a plain string."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if isinstance(runs, str):
        runs = [(runs, False)]
    for text, bold in runs:
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def make_table(headers, rows, widths_mm, body_size=9.0, status_col=None, band=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fixed_layout(t)
    hdr = t.rows[0]
    _repeat_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        _cell_text(c, h, body_size + 0.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_bg(c, HEADER_FILL)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            bold = (status_col is not None and i == status_col)
            _cell_text(cells[i], val, body_size, bold=bold)
            if band and ridx % 2 == 1:
                _set_cell_bg(cells[i], ALT_FILL)
    # widths on every row (fixed layout needs them set on cells)
    t.autofit = False
    t.allow_autofit = False
    for row in t.rows:
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================ CONTENT ========================================
add_title("Calibration — Solo Dry-Run Outcomes",
          "Prep-proposal vs result vs still-open  ·  FIN-9 + credential-confidence calibration")

add_meta("Date:", "2026-06-03      Status: Sandbox — nothing shipped (no master merge, no push, no deploy; all edits uncommitted).")
add_meta("Companion to:", "00_CALIBRATION_SUMMARY.md (the input brief). This sheet is the output — what the autonomous pass produced against the fixed-code baseline (test_fixtures/phishield_live.json).")
add_meta("Gate:", "verify_supply_chain_financial_wiring.py 31/31  ·  verify_scan_smoke.py exit 0 (59.6s, example.com 145/Low).")
add_meta("Legend:", "Locked = data-anchored, applied, sign-off only.   Gated = colleague-gated value (best-effort set).   Structural = design decision, not a parameter.")

# ---- Section 1: decision table ----
add_h2("1.  Decision table")
dec_headers = ["#", "Item", "Prep proposal (range)", "Dry-run result (applied)", "Status"]
dec_rows = [
    ["1", "Vuln curve shape",
     "fix polarity → convex (score/1000)^k, k 1.5–2.0",
     "polarity fixed (1cc204d); convex k=1.8. phishield vuln 0.169 → 0.0386",
     "Locked"],
    ["2", "Base rate (LEF 0.3)",
     "retain 0.20–0.35 if convex",
     "retained 0.3. p_breach 0.0735 → 0.0168",
     "Gated — highest leverage"],
    ["3", "Credential K1–K7 + caps",
     "K1 1.0/0.4/0.1, K2 decay, K3 ×0.3, dehashed→class, darkweb −40 / paste −30",
     "full K-model applied. phishield HIGH → LOW (W=0.432, contrib 10); takealot CRITICAL (infostealer floor, contrib 100)",
     "Locked; K3 combo-recency Gated"],
    ["4", "RSI factor rebalance",
     "RDP 0.18–0.22, critical-cred ≥ RDP, trim surfaces",
     "applied. phishield RSI 0.451 → 0.219 (dropped false HIGH-cred factor, 3 → 2); takealot critical-cred sole factor",
     "Locked"],
    ["5", "SA per-record / industry cost",
     "hold (IBM-2025-SA), add refresh stamp",
     "values held + dated sourcing stamp",
     "Locked"],
    ["6", "POPIA C2 fine",
     "P(fine)×E[fine] ≈ R100k–250k; hold R10m cat ceiling",
     "0.03 × R5m = R150k (was R200k @ R10m); R10m ceiling held for cat view",
     "Gated — P(fine|breach)"],
    ["7", "TEF SA tilt",
     "Gov 1.40–1.50, Comms 1.20–1.30",
     "Gov/Public Sector 1.45, Comms 1.25, Consumer 1.10; FS held 1.45",
     "Locked; Healthcare 1.40 trim = soft-open"],
    ["8", "FIN-9 Pareto LGB tail",
     "alpha 1.5–2.0, mix_w 0.25–0.35, f_sc 0.12",
     "alpha 1.77, mix_w 0.30, f_sc 0.12; severity-only on SC slice; supply_chain_tail_adjustment.applied=False kept",
     "Gated — alpha + mix_w"],
    ["9", "Risk-level bands",
     "re-fit to de-inflated dist + p_breach tiers",
     "Retain 200/400/600 (= inverse of neutral {2%,6%,12%}); comment-only, byte-identical",
     "Locked"],
]
make_table(dec_headers, dec_rows, [7, 25, 51, 61, 29], body_size=8.5, status_col=4)

# ---- Section 2: evidence ----
add_h2("2.  Evidence (verified from saved iterations)")
add_h3("phishield — clean FSP, R10m floor, Financial Services (iter0 → iter7)")
tr_headers = ["Metric", "iter0 (fixed-code)", "iter7 (final)", "Δ"]
tr_rows = [
    ["risk_score / level", "169 / Low", "164 / Low", "flat (correct)"],
    ["vulnerability", "0.169 (linear)", "0.0386 (convex k=1.8)", "curve"],
    ["p_breach", "0.0735", "0.0168", "−77%"],
    ["ML annual loss", Rfull(1793092), Rfull(750247), "−58%"],
    ["1-in-250 (P99.6)", Rfull(7989625), Rfull(4822195), "−40%"],
    ["credential class", "HIGH", "LOW", "de-escalated"],
]
make_table(tr_headers, tr_rows, [44, 44, 44, 36], body_size=9.0, status_col=3)
add_para([("The dominant movements are de-inflation (removing false-positive inflation that had landed on "
           "well-postured orgs), not parameter tuning. Wiring confirmed: ", False),
          ("p_breach = vulnerability × TEF × 0.3 = 0.0386 × 1.45 × 0.3 = 0.0168.", True)], size=10)

add_h3("takealot — R20bn, Consumer (iter8)")
add_para([("risk 235 / Medium  ·  p_breach 0.0243  ·  ML R", False), (NB, False), ("124.0m  ·  1-in-250 R", False),
          (NB, False), ("577.0m  ·  credential CRITICAL (infostealer)  ·  C4 ransom R", False), (NB, False),
          ("25.7m.  ", False), ("Flags C1 liability = 0 (see §3, item 5).", True)], size=10)

add_h3("FIN-9 tail-only proof (iter5 → iter6, phishield)")
add_para([("Median / ML / p_breach byte-flat; 1-in-100 ", False), ("+20%", True), (", 1-in-200 ", False),
          ("+38%", True), (", 1-in-250 ", False), ("+45.5%", True),
          (" — inside the doc-05 18–50% design band, ordering preserved, no double-count.", False)], size=10)

add_h3("Severity anchor")
add_para([("At the R200m median-revenue pivot (revenue_scale = 1.0) the Financial Services severity reproduces "
           "the IBM-SA Financial Services anchor (≈ R", False), (NB, False),
          ("70.1m) to within ~4% — the median is pinned to real SA breach data, not invented. "
           "(Task-10 median-pivot recompute; not saved as its own iteration file.)", False)], size=10)

# ---- Section 3: agenda ----
add_h2("3.  Still-open — the session agenda (ranked)")
add_h3("Tier A — needs the colleague's judgement")
agenda_a = [
    ("Base rate (the 0.3 LEF constant). ", "Annual loss-event (~1–3% SME) vs material-incident moves the constant 3–5×; everything downstream keys off it. Lock this first."),
    ("FIN-9 tail: alpha (1.77) + mix_w (0.30). ", "The colleague's domain (international breach-cost / EVT). alpha 1.2–2.5; mix_w 0.25–0.35. f_sc=0.12 is data-anchored (IBM CoDB), not gated."),
    ("POPIA P(fine | private-SME breach) = 0.03. ", "Inferred from enforcement scarcity (0 private fines; both s109 fines R5m, public-sector). A compliance-officer call."),
    ("K3 combo-recency interaction. ", "Flat ×0.3 is wrong for the fresh+combo case (e.g. ALIEN TXTBASE 2024-12 in a combolist) — recency-aware?"),
]
for i, (lead, rest) in enumerate(agenda_a, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(lead); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(rest); r2.font.size = Pt(10)

add_h3("Tier B — structural (design decision, surfaced by the dry-run, not a parameter)")
agenda_b = [
    ("C1 liability — give it its own factor in the cat model, not the residual balance. ",
     "Today C1 = max(0, severity − C2 − C3 − C4 − C5) is a residual/plug. C3/BI is independently "
     "revenue-scaled, overruns the breach anchor, and floors C1 to 0 for big orgs (takealot C1=0; cost_components "
     "also omits C3, so the visible ≈ R 30.8m understates the MC-derived ML R 124.0m). The residual is only the "
     "symptom — a residual cannot carry its own tail, yet liability is the heaviest-tailed bucket in real cyber "
     "cat (class actions / regulatory cascade). Direction: in the CAT model (mc_c1, :2521) model C1 as an "
     "independent severity + tail; keep the residual in the central/point estimate (:2342 — a real central loss "
     "can legitimately floor a bucket). Side effects: removes the floor artefact AND restores a non-zero C1+C2 "
     "severity for FIN-9 (:2531) to widen. Candidate drivers: records × per-record liability (records already "
     "estimated, currently disclosure-only) or an independent lognormal/Pareto anchor; demote the IBM total to a "
     "coherence cap. Also surface C3 in cost_components (display fix)."),
    ("Annual-expected-loss runs warm (~3.5%) for a clean R200m org. ",
     "Driven by the 7 incident-scenario probabilities + BI scaling, which sit outside the FAIR p_breach parameters "
     "calibrated here. Single-breach point estimates are well-anchored; the annualised aggregation is where the "
     "warmth enters."),
]
for lead, rest in agenda_b:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    # continue numbering 5,6 visually via manual prefix (List Number restarts per call group)
    r1 = p.add_run(lead); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(rest); r2.font.size = Pt(10)

add_h3("Tier C — resolved, sign-off only")
add_para([("TEF SA tilts  ·  risk bands (retain)  ·  severity anchor  ·  C4 ransom (Sophos-aligned)  ·  "
           "credential de-escalation. Do not spend session time unless challenged.", False)], size=10)

# ---- Section 4: data map ----
add_h2("4.  Where the raw data lives")
add_para([("Per-iteration numbers: tooling/_calib_iterations/iter0…iter8.json   ·   "
           "per-edit rationale: tooling/_apply_item02…9.py docstrings   ·   "
           "per-topic sources: docs/calibration_prep/01…05_*.md   ·   "
           "in-code anchored comments: scoring_analytics.py (band-retain block, FIN-9 basis, POPIA/TEF).",
           False)], size=9, italic=True, color=RGBColor(0x44, 0x44, 0x44))

# core props
doc.core_properties.title = "Calibration - Solo Dry-Run Outcomes"
doc.core_properties.subject = "FIN-9 + credential-confidence calibration (sandbox dry-run)"

# ---- R10 third-person assertion: grep assembled body text for banned pronouns
banned = re.compile(r"\b(we|us|our|i|you|your|they|their)\b", re.IGNORECASE)
hits = []
for para in doc.paragraphs:
    for m in banned.finditer(para.text):
        hits.append((para.text[:60], m.group(0)))
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for m in banned.finditer(cell.text):
                hits.append((cell.text[:40], m.group(0)))
assert not hits, ("R10 first-person voice violation", hits)

doc.save(OUT)
print("OK wrote", OUT)
print("   R10 third-person grep: clean (0 banned-pronoun hits)")
print("   tables:", len(doc.tables), " paragraphs:", len(doc.paragraphs))
