"""Generic Markdown -> Phishield-branded .docx converter.

Reuses the rendering helpers from generate_outstanding_docx.py (headings,
tables, bullets, inline formatting, brand colours) so any internal design /
pre-read document gets the same house style as OUTSTANDING.docx.

Usage (from security_scanner/ directory):
    py -3 tooling/md_to_branded_docx.py --src docs/foo.md --out docs/foo.docx \
        --title "Document Title" --subtitle "Optional subtitle line"
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE))
import generate_outstanding_docx as gen  # noqa: E402


def build(src, out, title, subtitle=None, strip_links=False):
    md_text = Path(src).read_text(encoding="utf-8")
    if strip_links:
        # Collapse [display](url) -> display. Keeps inline-code display
        # (`code`) intact (rendered in Consolas downstream) and drops noisy
        # repo file#line paths that an external reader does not need.
        md_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", md_text)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Cover block (Phishield house style)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PHISHIELD CYBER RISK SCANNER")
    r.bold = True
    r.font.color.rgb = gen.NAVY
    r.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(title)
    r.bold = True
    r.font.color.rgb = gen.NAVY
    r.font.size = Pt(19)

    if subtitle:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p3.add_run(subtitle)
        r.italic = True
        r.font.color.rgb = gen.GREY_MID
        r.font.size = Pt(9.5)

    doc.add_paragraph()

    # The source .md repeats the title as its first H1; skip that one so it
    # isn't duplicated under the cover.
    skipped_first_h1 = False
    # Buffer consecutive 'p' lines into ONE paragraph (correct Markdown
    # semantics: a paragraph runs until a blank line). The shared line-based
    # parser emits one op per physical line, which would otherwise break inline
    # spans (e.g. **bold**) that are hard-wrapped across source lines.
    para_buf = []

    def flush_para():
        if para_buf:
            gen._apply_inline_formatting(doc.add_paragraph(), " ".join(para_buf))
            para_buf.clear()

    for op in gen._parse_md(md_text):
        kind = op[0]
        if kind == "p":
            para_buf.append(op[1].strip())
            continue
        flush_para()
        if kind == "h1":
            if not skipped_first_h1:
                skipped_first_h1 = True
                continue
            gen._add_heading(doc, op[1], 1)
        elif kind == "h2":
            gen._add_heading(doc, op[1], 2)
        elif kind == "h3":
            gen._add_heading(doc, op[1], 3)
        elif kind == "bullet":
            gen._apply_inline_formatting(
                doc.add_paragraph(style="List Bullet"), op[1])
        elif kind == "table":
            gen._add_table(doc, op[1])
            doc.add_paragraph()
        elif kind == "hr":
            hp = doc.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr = hp.add_run("— — —")
            hr.font.color.rgb = gen.GREY_MID
            hr.font.size = Pt(9)
    flush_para()

    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Generated: {out}")
    print(f"Size: {Path(out).stat().st_size // 1024} KB")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--title", required=True)
    ap.add_argument("--subtitle", default=None)
    ap.add_argument("--strip-links", action="store_true",
                    help="Collapse [text](url) to text (drop repo paths)")
    a = ap.parse_args()
    build(a.src, a.out, a.title, a.subtitle, strip_links=a.strip_links)
