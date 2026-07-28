"""Generate the Phishield Scanner DEMO GUIDE (.docx) — the tester-facing manual.

Run: py generate_demo_guide.py

WHY A SEPARATE DOCUMENT
    The full user manual (`Phishield_Cyber_Risk_Scanner_User_Manual.docx`) is an
    internal document. It documents the scoring weights, calibration anchors,
    financial-model parameters, per-checker detection logic and the third-party
    data sources behind each card. That is the product. Handing it to an external
    tester hands over a specification someone could build against.

    This guide gives a demo user everything needed to RUN the scanner and READ its
    output, and nothing that would let them reproduce it.

DISCLOSURE RULE (apply to every edit of this file)
    Describe WHAT a signal means for the client's risk and WHAT to do about it.
    Never describe HOW the signal is obtained or weighted. Concretely, this file
    must not name:
      * any third-party data provider, feed or API the scanner queries;
      * any scoring weight, multiplier, band, threshold or calibration anchor;
      * the financial model's distributions, parameters or severity mapping;
      * per-checker detection logic, evidence gates or false-positive controls;
      * internal identifiers (SCN-*, checker module names, env vars).
    There is a test at the bottom of this module that enforces the naming half of
    that rule mechanically, so a careless future edit fails the build.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from manual_parts import helpers as H

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      "Phishield_Scanner_Demo_Guide.docx")

# Terms that must never appear in the generated document. Provider names, model
# internals and tuning knobs are the replicable part of the product.
FORBIDDEN = [
    # data sources / vendors
    "dehashed", "intelx", "intelligence x", "have i been pwned", "hibp", "shodan",
    "virustotal", "securitytrails", "crt.sh", "certspotter", "ransomware.live",
    "hudson rock", "gemini", "perplexity", "veilguard", "censys", "greynoise",
    # model internals / tuning
    "monte carlo", "pert", "lognormal", "poisson", "weight", "multiplier",
    "calibration", "anchor", "threshold", "rsi", "ransomware susceptibility",
    "return period", "1-in-250", "p95", "p50", "severity band", "uplift",
    # implementation
    "checker", "scoring_analytics", "scanner.py", "env var", "api key",
    "tarpit", "canary", "evidence gate", "cve gating", "scn-",
]


def _cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PhiShield CyberRisk Scanner")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = H.NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Demo Guide")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("How to run a scan and read the results")
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(110, 110, 110)

    for _ in range(9):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Evaluation release. Confidential.\n"
                  "Prepared for named demo participants only; please do not "
                  "forward this document or your login.")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(130, 130, 130)


def _s1_what_it_is(doc):
    H.add_h1(doc, "1. What the scanner does")
    H.add_body(doc,
        "The PhiShield CyberRisk Scanner assesses an organisation's cyber risk from "
        "the outside, the same way an attacker first sees it. You give it a domain "
        "name. It examines what that organisation exposes to the internet, what has "
        "already leaked about it, and how that combination translates into insurable "
        "loss, then returns a risk score, a prioritised list of what to fix, and a "
        "financial view of what an incident could cost.")
    H.add_body(doc,
        "It is designed for underwriting and broking conversations. The output is "
        "meant to be defensible in front of a client, so every finding is evidence "
        "based: the scanner reports what it can actually observe, and says so "
        "plainly when it cannot observe something.")

    H.add_h2(doc, "1.1 What it is not")
    H.add_bullet(doc,
        "It is not a penetration test. Nothing is exploited, no credentials are "
        "guessed, no vulnerability is triggered.")
    H.add_bullet(doc,
        "It is not intrusive and needs no cooperation from the target. There is "
        "nothing to install, no agent, no access, no scheduled window.")
    H.add_bullet(doc,
        "It is not continuous. A scan is a snapshot of the moment it ran. Exposure "
        "changes, so a result ages.")
    H.add_bullet(doc,
        "It is not a compliance certification. It informs a risk view; it does not "
        "certify a standard.")
    H.add_note(doc,
        "Because the scan is external and passive, you can run it on a prospect "
        "before any relationship exists. That is the point: it is a conversation "
        "opener, not an audit.")


def _s2_access(doc):
    H.add_h1(doc, "2. Getting in")
    H.add_body(doc,
        "The demo runs at the address below. Your browser will ask for a username "
        "and password the first time you visit. Use the credentials issued to you "
        "individually; they are not shared, so please do not pass them on.")
    p = doc.add_paragraph()
    r = p.add_run("https://veilguard.phishield.com/scanner/")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = H.NAVY
    H.add_body(doc,
        "If the login is rejected, check that the password was copied without a "
        "trailing space. If it is still rejected, your access may have been ended; "
        "contact your PhiShield representative.")
    H.add_warning(doc,
        "This is an evaluation environment holding real data about real "
        "organisations. Treat everything you see, and everything you download, as "
        "confidential. Do not scan a domain you have no legitimate interest in.")


def _s3_running(doc):
    H.add_h1(doc, "3. Running a scan")
    H.add_body(doc,
        "Enter the organisation's primary domain, for example acme.co.za. Do not "
        "include http:// or a path. Then supply two pieces of context:")
    H.add_bold_body(doc, "Industry. ",
        "Loss patterns differ sharply by sector, so this shapes both the likelihood "
        "and the cost side of the analysis. Choose the closest match to what the "
        "organisation actually does, not how it describes itself.")
    H.add_bold_body(doc, "Annual revenue (ZAR). ",
        "Scale drives exposure: the same weakness costs a very different amount at "
        "different sizes. A reasonable estimate is fine, but a wrong order of "
        "magnitude will make the financial section wrong. If you do not know, use "
        "the best public figure rather than leaving it at a default.")
    H.add_body(doc,
        "A scan usually completes in a few minutes. Progress is shown live. You can "
        "leave the page and come back; the result is saved against the domain.")
    H.add_tip(doc,
        "Scan the domain the organisation actually trades under. Holding-company "
        "domains often have almost no external footprint and will look misleadingly "
        "clean.")


def _s4_reading(doc):
    H.add_h1(doc, "4. Reading the result")

    H.add_h2(doc, "4.1 The headline score")
    H.add_body(doc,
        "The overall score summarises the organisation's external risk position, "
        "with a plain-language risk level beside it. Higher means more risk. It is a "
        "composite: a single severe weakness can lift it sharply even when most of "
        "the estate looks healthy, which is deliberate, because that is how losses "
        "actually happen.")
    H.add_note(doc,
        "Treat the score as a conversation starter and a way to compare like with "
        "like over time. The findings underneath it are what you act on.")

    H.add_h2(doc, "4.2 The finding cards")
    H.add_body(doc,
        "Findings are grouped into themes covering, broadly: what the organisation "
        "exposes to the internet, how its email domain is protected against "
        "impersonation, the health of its public-facing systems, what has already "
        "leaked about its people and accounts, and its posture against the failure "
        "modes insurers care about. Each card states what was observed, why it "
        "matters commercially, and what fixing it involves.")
    H.add_body(doc,
        "Click a card to open the supporting evidence. Anything asserted in a report "
        "can be traced to something observed.")

    H.add_h2(doc, "4.3 When something says 'not assessed'")
    H.add_body(doc,
        "Occasionally a card reports that an area could not be assessed. This is not "
        "the same as a clean result, and the scanner will never present it as one. "
        "It means the evidence needed to reach a conclusion was unavailable at scan "
        "time. Treat it as an open question to revisit, not as a pass.")

    H.add_h2(doc, "4.4 The remediation queue")
    H.add_body(doc,
        "This is the ordered answer to 'what should they do first'. It is sequenced "
        "by the reduction in risk each action buys relative to the effort it takes, "
        "so working from the top down is the fastest route to a materially better "
        "position. For a renewal conversation, the first three items are usually the "
        "whole discussion.")

    H.add_h2(doc, "4.5 The financial view")
    H.add_body(doc,
        "The financial section converts the security picture into money: the range "
        "of loss the organisation might reasonably suffer, and how bad a genuinely "
        "severe year could get. It is expressed as a range rather than a single "
        "number because the honest answer is a distribution, and a point estimate "
        "would imply precision that does not exist.")
    H.add_body(doc,
        "Use the middle of the range for ordinary planning conversations and the "
        "severe end for limit adequacy. The numbers are modelled estimates for "
        "discussion, not a quotation, a reserve, or a promise of cover.")


def _s5_reports(doc):
    H.add_h1(doc, "5. The reports")
    H.add_body(doc,
        "Three PDFs are produced from a single scan, each written for a different "
        "reader. They draw on the same findings; they differ in depth, not in "
        "substance.")
    H.add_bold_body(doc, "Executive summary. ",
        "For the client's board or owner. The position, the money, and the few "
        "things that matter, with no jargon.")
    H.add_bold_body(doc, "Broker report. ",
        "For the placing conversation. Adds the prioritised remediation view and "
        "enough substantiation to defend the position with an underwriter.")
    H.add_bold_body(doc, "Full technical report. ",
        "For the client's IT function or provider. Every finding with its "
        "supporting evidence and remediation detail.")
    H.add_tip(doc,
        "Send the executive summary first. If the client asks 'how do you know "
        "that', the full technical report is the answer, and it is better received "
        "when it arrives on request.")


def _s6_limits(doc):
    H.add_h1(doc, "6. Limitations worth stating out loud")
    H.add_body(doc,
        "Being straight about these protects the credibility of everything else in "
        "the report.")
    H.add_bullet(doc,
        "A scan reflects a moment in time. Something fixed the next morning still "
        "appears in yesterday's report.")
    H.add_bullet(doc,
        "Only the externally visible estate is assessed. Internal controls, backups, "
        "staff training and incident response are invisible from outside and are not "
        "judged here. A clean external result does not mean a well-run organisation.")
    H.add_bullet(doc,
        "Absence of evidence is not evidence of absence. Nothing found does not "
        "prove nothing happened.")
    H.add_bullet(doc,
        "The financial figures are modelled estimates for discussion. They are not "
        "an offer, a quotation, or confirmation that any loss would be covered.")
    H.add_bullet(doc,
        "Shared and outsourced infrastructure can behave unusually. Where the "
        "scanner is uncertain who operates something, it prefers to say so rather "
        "than attribute it to the client.")


def _s7_data(doc):
    H.add_h1(doc, "7. Handling what you find")
    H.add_body(doc,
        "Scan results describe real organisations and, in places, real people. Under "
        "POPIA that carries obligations on you as well as on us.")
    H.add_bullet(doc,
        "Share a report only with the organisation it concerns, or internally on a "
        "need-to-know basis.")
    H.add_bullet(doc,
        "Where exposed accounts are listed, identifying detail is deliberately "
        "masked in every report and on screen. Do not attempt to unmask it.")
    H.add_bullet(doc,
        "Never test a leaked credential against any live system. That is unlawful, "
        "regardless of intent.")
    H.add_bullet(doc,
        "Delete downloaded reports when the evaluation ends.")

    H.add_h2(doc, "7.1 Credential release is switched off for this demo")
    H.add_body(doc,
        "In normal operation the scanner can release the full detail of an "
        "organisation's exposed credentials to that organisation, encrypted so that "
        "only they can open it. That release is available only against the client's "
        "signed consent, recorded at the time of the request.")
    H.add_body(doc,
        "The page that captures and records that consent is not part of this "
        "evaluation, so the release itself is disabled. You will see the control "
        "marked unavailable. Everything else about the exposure, which accounts are "
        "affected, where they were exposed and how recently, is fully present in the "
        "reports. Nothing about the risk assessment is diminished by this.")


def _s8_feedback(doc):
    H.add_h1(doc, "8. What we would like from you")
    H.add_body(doc,
        "The most useful feedback is specific and tied to a scan you actually ran.")
    H.add_bullet(doc,
        "Anything that looks wrong. If a finding does not match what you know about "
        "an organisation, tell us the domain and the card. A confirmed false "
        "positive is the single most valuable thing you can send us.")
    H.add_bullet(doc,
        "Anything you expected to see and did not.")
    H.add_bullet(doc,
        "Whether the executive summary would survive being put in front of your "
        "client unedited, and if not, what you would change.")
    H.add_bullet(doc,
        "Where the wording sends the wrong message, is too alarming, or is too soft.")
    H.add_body(doc,
        "Please include the domain and roughly when you ran it, so the exact result "
        "can be retrieved.")
    H.add_note(doc,
        "Access is time limited for the evaluation period and will be withdrawn "
        "afterwards. Anything you want to keep should be saved before then, subject "
        "to section 7.")


def build(doc):
    _cover(doc)
    doc.add_page_break()
    for i, section in enumerate((_s1_what_it_is, _s2_access, _s3_running, _s4_reading,
                                 _s5_reports, _s6_limits, _s7_data, _s8_feedback)):
        if i:
            doc.add_paragraph()
        section(doc)


def _assert_no_disclosure(doc):
    """Fail the build if anything on the do-not-disclose list reached the page.

    A prose rule in a docstring decays; this does not. It is a naming check, not a
    semantic one, so it cannot prove the document is free of methodology, but it
    does catch the realistic failure: pasting a paragraph across from the internal
    manual.
    """
    text = "\n".join(p.text for p in doc.paragraphs).lower()
    # The public demo host is exempt: testers have to type it, and it is in their
    # browser bar regardless. Masking it rather than dropping "veilguard" from the
    # list keeps a stray reference to the internal stack a build failure.
    text = text.replace("veilguard.phishield.com", "<demo-host>")
    hits = sorted({t for t in FORBIDDEN if t in text})
    if hits:
        raise AssertionError(
            "Demo guide discloses internal terms: " + ", ".join(hits) +
            "\nThe demo guide must describe what findings MEAN, never how they are "
            "obtained or weighted. See the disclosure rule at the top of this file.")


def main():
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(10)
    build(doc)
    _assert_no_disclosure(doc)
    doc.save(OUTPUT)
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Demo guide saved to: {OUTPUT}")
    print(f"   {len(doc.paragraphs)} paragraphs, ~{words} words")
    print("   disclosure check: PASS (no internal terms present)")


if __name__ == "__main__":
    main()
